import os
import streamlit as st
import time
from langchain_groq import ChatGroq
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain.chains.combine_documents import create_stuff_documents_chain
from langchain_core.prompts import ChatPromptTemplate
from langchain.chains import create_retrieval_chain
from langchain_community.vectorstores import FAISS
from langchain_community.document_loaders import PyPDFLoader
from langchain_google_genai import GoogleGenerativeAIEmbeddings
from docx import Document
from langchain.schema import Document as LangchainDocument
import warnings
import shutil

def main_document():

    # Load the GROQ And Google API KEY
    groq_api_key = st.secrets['GROQ_API_KEY']
    os.environ["google_api_key"] = st.secrets['GOOGLE_API_KEY']

    # Add the logo to the sidebar
    st.title("Ask a question!")

    models = {
        "Llama3-70B": "llama3-70b-8192",
        "Llama3-8B": "llama3-8b-8192",
        "Gemma-7B": "gemma-7b-it",
        "Mixtral-8x": "Mixtral-8x7b-32768"
    }

    selected_model = st.sidebar.selectbox(
        "**Select Model**",
        list(models.keys())
    )

    llm = ChatGroq(groq_api_key=groq_api_key,
                   model_name=models[selected_model])

    prompt_temp = ChatPromptTemplate.from_template(
        """
        Answer the questions based on the provided context only.
        Please provide the most accurate response based on the question
        <context>
        {context}
        <context>
        Questions:{input}
        """
    )

    # Create a temporary directory to save uploaded files
    temp_dir = "temp_files"
    if not os.path.exists(temp_dir):
        os.makedirs(temp_dir)

    def load_docx(file_path):
        """Function to load and extract text from a DOCX file"""
        doc = Document(file_path)
        full_text = []
        for para in doc.paragraphs:
            full_text.append(para.text)
        return '\n'.join(full_text)

    def vector_embedding(uploaded_files):
        if "vectors" not in st.session_state:
            st.session_state.embeddings = GoogleGenerativeAIEmbeddings(model="models/embedding-001")
            documents = []

            # Process each uploaded file
            for uploaded_file in uploaded_files:
                # Save the uploaded file temporarily
                temp_file_path = os.path.join(temp_dir, uploaded_file.name)
                with open(temp_file_path, "wb") as f:
                    f.write(uploaded_file.getbuffer())

                # Load the file content based on the file type
                if uploaded_file.name.endswith('.pdf'):
                    loader = PyPDFLoader(temp_file_path)
                    loaded_docs = loader.load()
                    for doc in loaded_docs:
                        documents.append(LangchainDocument(page_content=doc.page_content, metadata=doc.metadata))
                elif uploaded_file.name.endswith('.docx'):
                    docx_text = load_docx(temp_file_path)
                    documents.append(LangchainDocument(page_content=docx_text, metadata={'source': temp_file_path}))

            st.session_state.text_splitter = RecursiveCharacterTextSplitter(chunk_size=1000, chunk_overlap=200)
            st.session_state.final_documents = st.session_state.text_splitter.split_documents(documents)
            st.session_state.vectors = FAISS.from_documents(st.session_state.final_documents, st.session_state.embeddings)

            # Delete the temporary files after vector creation
            for uploaded_file in uploaded_files:
                temp_file_path = os.path.join(temp_dir, uploaded_file.name)
                os.remove(temp_file_path)

    uploaded_files = st.sidebar.file_uploader("Upload PDF or DOCX files", type=["pdf", "docx"], accept_multiple_files=True)

    # Disable the button if no files are uploaded
    if len(uploaded_files) > 0:
        upload_button_disabled = False
    else:
        upload_button_disabled = True

    if st.sidebar.button("Upload & Process Files", disabled=upload_button_disabled):
        if uploaded_files:
            vector_embedding(uploaded_files)
            st.write(
                "Your uploaded files are converted into Vector Store DB. Now it is ready. You can ask any question from the uploaded files."
            )
        else:
            st.error("Please upload PDF or DOCX files first.")

    isPromptAvailable = False

    prompt = st.text_input("", placeholder="Enter a Prompt Here", key="user_input")

    if prompt:
        if "vectors" in st.session_state:
            document_chain = create_stuff_documents_chain(llm, prompt_temp)
            retriever = st.session_state.vectors.as_retriever()
            retrieval_chain = create_retrieval_chain(retriever, document_chain)
            start = time.process_time()
            response = retrieval_chain.invoke({'input': prompt})
            st.divider()
            with st.container():
                col1, col2 = st.columns(2)
                with col1:
                    st.write(f"Response time: {time.process_time() - start}")
                with col2:
                    st.write(f"Selected Model: {selected_model}")
                st.success(response['answer'])

                # With a streamlit expander
                with st.expander("Document Similarity Search"):
                    # Find the relevant chunks
                    for i, doc in enumerate(response["context"]):
                        st.write(doc.page_content)
                        st.write("--------------------------------")

            isPromptAvailable = True
        else:
            st.error("Please upload and process the documents first! open sidebar! 👈")

    # Help button on the top-right of the main app
    st.markdown(
        """
        <style>
        .help-button {
            position: absolute;
            top: 10px;
            right: 10px;
            background-color: #4CAF50;
            color: white;
            padding: 10px 20px;
            border: none;
            border-radius: 4px;
            cursor: pointer;
        }
        </style>
        <button class="help-button" onclick="document.getElementById('help-modal').style.display='block'">Help</button>
        <div id="help-modal" style="display:none;">
            <div style="position:fixed;top:50%;left:50%;transform:translate(-50%,-50%);background-color:white;padding:20px;border-radius:10px;box-shadow:0px 0px 10px rgba(0,0,0,0.5);z-index:1000;">
                <h3>Instructions</h3>
                <p>This application allows you to upload and analyze documents for quick information retrieval.</p>
                <ul>
                    <li>1. Upload your PDF or DOCX files using the uploader below.</li>
                    <li>2. Click the 'Upload & Process Files' button to process your files.</li>
                    <li>3. Once processed, you can enter any question related to the content of the uploaded documents, and the AI will provide an accurate response.</li>
                </ul>
                <button onclick="document.getElementById('help-modal').style.display='none'">Close</button>
            </div>
        </div>
        """,
        unsafe_allow_html=True
    )

    st.sidebar.markdown('© 2024 Brainwave AI.')

    # Cleanup temporary directory on session state initialization
    def cleanup_temp_dir():
        if os.path.exists(temp_dir):
            shutil.rmtree(temp_dir)
            os.makedirs(temp_dir)

    if "initialized" not in st.session_state:
        st.session_state.initialized = True
        cleanup_temp_dir()

    # Custom CSS for sidebar menu
    hide_default_loader = """
        <style>
        div[data-testid="stToolbar"] {{    
            visibility: hidden;    
            height: 50%;    
            position: fixed;
        }}
        div[data-testid="stDecoration"]{{    
            visibility: hidden;    
            height: 0%;    
            position: fixed;
        }}
        .stFileUploader{{    
            margin-top : 4%
        }}
        div[data-testid="stStatusWidget"]{{    
            visibility: hidden;    
            height: 50%;    
            position: fixed;
        }}
        #MainMenu{{    
            visibility: hidden;    
            height: 0%;
        }}
        header{{    
            visibility: hidden;    
            height: 0%;
        }}
        footer{{   
            visibility: hidden;
            height: 0%;
        }}
        .st-emotion-cache-qcqlej{{
            display:{"block" if (prompt) else "none"};
            flex-grow:{"1" if (prompt) else ""}; 
        }}
        .st-emotion-cache-bm2z3a{{
            justify-content:{"" if (prompt) else "center"}; 
        }}
        .st-emotion-cache-1eo1tir{{
            padding:6rem
        }}
        </style>
    """
    st.markdown(hide_default_loader, unsafe_allow_html=True)


if __name__ == "__main__":
    main_document()
